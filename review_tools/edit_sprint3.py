from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(
    "C:/Users/bruda/.codex/visualizations/2026/09/01/"
    "01a05dd2-e345-7362-a8fb-3c96f1a73cff/document-audit-session/"
    "artifacts/sources/sprint3.docx"
)
OUTPUT = Path(
    "C:/Users/bruda/.codex/visualizations/2026/09/01/"
    "01a05dd2-e345-7362-a8fb-3c96f1a73cff/document-audit-session/final/"
    "Documentacion_Sprint_3_Integracion_Multiplayer_Experimentacion_y_Testing_ACTUALIZADA_PLAN_QA.docx"
)


def set_paragraph_text(paragraph, text):
    """Replace text while preserving the first run's formatting and paragraph XML."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_labeled_paragraph(paragraph, label, body):
    """Preserve the document's existing label/body run formatting."""
    if not paragraph.runs:
        paragraph.add_run(label)
        paragraph.add_run(body)
        return

    prototypes = list(paragraph.runs)
    paragraph.runs[0].text = label

    if len(paragraph.runs) >= 2:
        paragraph.runs[1].text = body
    else:
        new_run = paragraph.add_run(body)
        if prototypes[0]._r.rPr is not None:
            new_run._r.insert(0, deepcopy(prototypes[0]._r.rPr))
        new_run.bold = False
        new_run.font.color.rgb = RGBColor(0x24, 0x31, 0x3D)
        new_run.font.size = Pt(9)

    for run in paragraph.runs[2:]:
        run.text = ""


def set_cell_text(cell, text):
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def cell(table, row, column, text):
    set_cell_text(table.rows[row].cells[column], text)


def remove_page_break(paragraph):
    for element in list(paragraph._p.iter(qn("w:br"))):
        if element.get(qn("w:type")) == "page":
            element.getparent().remove(element)


def prevent_row_split(row):
    """Keep one table row together when Word paginates the document."""
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def iter_all_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for table_cell in row.cells:
                for paragraph in table_cell.paragraphs:
                    yield paragraph


def restore_spanish_diacritics(document):
    word_replacements = {
        "actualizacion": "actualización",
        "ademas": "además",
        "anonima": "anónima",
        "aprobacion": "aprobación",
        "automaticamente": "automáticamente",
        "automaticos": "automáticos",
        "automatizados": "automatizados",
        "camara": "cámara",
        "catalogo": "catálogo",
        "caida": "caída",
        "caidas": "caídas",
        "codigo": "código",
        "codigos": "códigos",
        "colision": "colisión",
        "colisiones": "colisiones",
        "compilacion": "compilación",
        "comunicacion": "comunicación",
        "configuracion": "configuración",
        "conexion": "conexión",
        "critico": "crítico",
        "criticos": "críticos",
        "desconexion": "desconexión",
        "decision": "decisión",
        "dieciseis": "dieciséis",
        "ejecucion": "ejecución",
        "estan": "están",
        "fisica": "física",
        "fotografia": "fotografía",
        "identificacion": "identificación",
        "implementacion": "implementación",
        "integracion": "integración",
        "interaccion": "interacción",
        "item": "ítem",
        "items": "ítems",
        "segun": "según",
        "simultanea": "simultánea",
        "seccion": "sección",
        "gestion": "gestión",
        "mayusculas": "mayúsculas",
        "encontro": "encontró",
        "despues": "después",
        "modulos": "módulos",
        "academico": "académico",
        "podia": "podía",
        "comparo": "comparó",
        "priorizo": "priorizó",
        "linea": "línea",
        "mecanica": "mecánica",
        "mecanicas": "mecánicas",
        "metricas": "métricas",
        "minimo": "mínimo",
        "observacion": "observación",
        "observaciones": "observaciones",
        "podian": "podían",
        "posicion": "posición",
        "posiciones": "posiciones",
        "reaparecia": "reaparecía",
        "regresion": "regresión",
        "resolucion": "resolución",
        "reproduccion": "reproducción",
        "revision": "revisión",
        "rotacion": "rotación",
        "sesion": "sesión",
        "sesiones": "sesiones",
        "sincronizacion": "sincronización",
        "tecnica": "técnica",
        "tecnicas": "técnicas",
        "tecnico": "técnico",
        "tecnicos": "técnicos",
        "unico": "único",
        "unica": "única",
        "validacion": "validación",
        "verificacion": "verificación",
        "confirmacion": "confirmación",
        "version": "versión",
        "versiones": "versiones",
    }
    for paragraph in iter_all_paragraphs(document):
        for run in paragraph.runs:
            value = run.text
            for plain, accented in word_replacements.items():
                def apply_case(match):
                    original = match.group(0)
                    if original[:1].isupper():
                        return accented[:1].upper() + accented[1:]
                    return accented

                value = re.sub(
                    rf"\b{plain}\b",
                    apply_case,
                    value,
                    flags=re.IGNORECASE,
                )
            run.text = value


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)

    if len(doc.paragraphs) != 133 or len(doc.tables) != 29:
        raise RuntimeError(
            f"Estructura inesperada: {len(doc.paragraphs)} parrafos, "
            f"{len(doc.tables)} tablas"
        )

    p = doc.paragraphs
    t = doc.tables

    # Portada y contexto.
    set_paragraph_text(p[0], "INFORME DE SPRINT")
    set_paragraph_text(
        p[5],
        "Base heredada del Sprint 2: Unity 6.3 LTS, C#, Rigidbody, escena y "
        "recorrido de parkour, camara, movimiento, salto, gravedad, colisiones "
        "y respawn. La configuracion integrada del Sprint 3 usa Unity "
        "6000.3.22f1, Netcode for GameObjects 2.13.2, Multiplayer Services/Relay "
        "2.3.1 y arquitectura Host/Client. Las PoC LE0-78 y LE0-80 permanecen "
        "como antecedentes tecnicos."
    )
    set_paragraph_text(
        p[8],
        "Al corte del 01/09/2026, Jira registra 27 items y 118 Story Points para "
        "el Sprint 3: 19 items (78 SP) finalizados, 7 items (38 SP) en curso o "
        "pruebas y 1 item (2 SP) por hacer. Los estados de implementacion se "
        "distinguen de la validacion manual: una issue finalizada no se registra "
        "automaticamente como PASS sin una ejecucion y evidencia identificables."
    )
    set_paragraph_text(
        p[10],
        "Cada seccion corresponde de forma directa a uno de los dieciseis "
        "entregables oficiales del Sprint 3. Jira es la fuente operativa; este "
        "informe presenta la fotografia del 01/09/2026 y separa configuracion "
        "inspeccionada, compilacion correcta, ejecucion de pruebas y evidencia."
    )

    # 1. Integracion multiplayer.
    cell(t[3], 1, 1, "Unity 6000.3.22f1; NGO 2.13.2; Multiplayer Services/Relay 2.3.1; Unity Test Framework 1.6.0.")
    cell(t[3], 2, 1, "Host/Client con autenticacion anonima. El Host crea una allocation de Relay y comparte un join code; el cliente usa ese codigo para incorporarse.")
    cell(t[3], 3, 1, "GameScene - Assets/Scenes/GameScene.unity, cargada desde MenuScene mediante NetworkSceneManager.")
    cell(t[3], 4, 1, "NetworkManager, UnityTransport y LobbyPlayerSpawner en MenuScene; prefab de jugador Assets/Personajes-objetos/Jugador.prefab con NetworkObject y NetworkTransform.")
    cell(t[3], 5, 1, "Rama feat/adicion-de-elementos-en-partida; commit evaluado c33e5c4. Compilacion C# validada; build ejecutable para dos clientes pendiente.")
    cell(t[3], 6, 1, "Corte: 01/09/2026. Responsabilidad compartida por LCS Equipo 07; integracion multiplayer a cargo de Tomas Molina Varas.")

    cell(t[4], 1, 1, "IMPLEMENTADO / COMPILA. Validacion manual con dos clientes pendiente.")
    cell(t[4], 2, 1, "IMPLEMENTADO. Host y cliente usan GameScene; ejecucion conjunta pendiente.")
    cell(t[4], 3, 1, "IMPLEMENTADO. Movimiento, salto, wall running, gancho y respawn estan integrados; regresion manual pendiente.")
    cell(t[4], 4, 1, "IDENTIFICADO. Unity 6000.3.22f1; rama y commit c33e5c4 registrados.")
    set_labeled_paragraph(
        p[24],
        "Evidencia: ",
        "el checkout c33e5c4 compila sin errores. La hoja QA reporta PASS para "
        "TST-S3-003/004 en 61dd628, no disponible localmente. Sin sesion manual "
        "documentada con dos clientes."
    )

    # 2. Crear e ingresar a partida.
    set_paragraph_text(
        p[27],
        "El jugador Host selecciona CREATE GAME, inicializa Unity Services y la "
        "autenticacion anonima, crea una allocation de Relay, obtiene un join code "
        "y arranca el Host. El segundo jugador ingresa el codigo, obtiene la "
        "JoinAllocation y arranca como cliente. Solo el Host habilita el inicio y "
        "carga GameScene mediante la gestion de escenas de NGO."
    )
    cell(t[5], 1, 1, "CREATE GAME -> inicializar servicios -> crear allocation Relay -> StartHost. El codigo se muestra y se copia al portapapeles.")
    cell(t[5], 2, 1, "Join code de Relay normalizado a mayusculas.")
    cell(t[5], 3, 1, "Ingresar join code -> JoinAllocationAsync -> configurar UnityTransport con DTLS -> StartClient.")
    cell(t[5], 4, 1, "GameScene, cargada por el Host mediante NetworkSceneManager.LoadScene.")
    cell(t[5], 5, 1, "Se valida codigo vacio y se capturan errores de servicios, allocation, conexion y desconexion con mensajes de estado. Prueba manual de errores pendiente.")

    cell(t[6], 1, 1, "IMPLEMENTADO / Jira Finalizada. Validacion manual pendiente.")
    cell(t[6], 2, 1, "IMPLEMENTADO. Relay conserva la sesion mientras el Host esta activo; validacion manual pendiente.")
    cell(t[6], 3, 1, "IMPLEMENTADO / Jira Finalizada. Ingreso real de un segundo jugador pendiente de ejecucion documentada.")
    cell(t[6], 4, 1, "SIN EVIDENCIA DE EJECUCION. No se afirma PASS hasta probar codigos validos, invalidos y desconexion.")

    # 3. Dos jugadores simultaneos.
    cell(t[7], 1, 0, "REV-01 / 01-09-2026")
    cell(t[7], 1, 1, "Host configurado")
    cell(t[7], 1, 2, "Cliente configurado")
    cell(t[7], 1, 3, "Relay / DTLS")
    cell(t[7], 1, 4, "REVISION TECNICA: COMPILA")
    cell(t[7], 2, 0, "TST-MP-01 / pendiente")
    cell(t[7], 2, 1, "Editor o build")
    cell(t[7], 2, 2, "Segunda build")
    cell(t[7], 2, 3, "Relay")
    cell(t[7], 2, 4, "NO EJECUTADA")
    set_labeled_paragraph(
        p[37],
        "Tiempo minimo observado y evidencia: ",
        "no se registró una sesion temporizada con dos instancias. La prueba manual "
        "TST-MP-01 queda pendiente y no se asigna PASS/FAIL/BLOCKED."
    )

    # 4. Mismo escenario y spawn.
    cell(t[8], 1, 1, "GameScene - Assets/Scenes/GameScene.unity.")
    cell(t[8], 1, 2, "IMPLEMENTADO; validacion manual pendiente.")
    cell(t[8], 2, 1, "Objeto Respawn; posicion centrada segun el indice del cliente.")
    cell(t[8], 2, 2, "IMPLEMENTADO; validacion manual pendiente.")
    cell(t[8], 3, 1, "Mismo Respawn con offset lateral calculado por LobbyPlayerSpawner.")
    cell(t[8], 3, 2, "IMPLEMENTADO; validacion manual pendiente.")
    cell(t[8], 4, 1, "Separacion de 1,5 unidades entre jugadores conectados.")
    cell(t[8], 4, 2, "IMPLEMENTADO; validacion manual pendiente.")
    cell(t[8], 5, 1, "SpawnAsPlayerObject(clientId, true) asigna ownership al cliente correspondiente.")
    cell(t[8], 5, 2, "IMPLEMENTADO; validacion manual pendiente.")
    cell(t[8], 6, 1, "El servidor valida la zona de muerte; el propietario aplica el respawn y NetworkTransform.Teleport evita interpolar el recorrido.")
    cell(t[8], 6, 2, "IMPLEMENTADO tras correcciones c33e5c4/99f19f3; validacion con cliente remoto pendiente.")

    # 5. Representacion remota.
    cell(t[9], 1, 1, "IMPLEMENTADO. Validacion visual con dos clientes pendiente.")
    cell(t[9], 2, 1, "IMPLEMENTADO mediante NetworkTransform; prueba de fluidez pendiente.")
    cell(t[9], 3, 1, "IMPLEMENTADO. PlayerNetworkSetup habilita camara y AudioListener solo para IsOwner.")
    cell(t[9], 4, 1, "IMPLEMENTADO. Move, CameraMovement, WallRun y Grappling solo se habilitan para el propietario.")
    cell(t[9], 5, 1, "Assets/Personajes-objetos/Jugador.prefab; rama feat/adicion-de-elementos-en-partida; commit c33e5c4.")

    # 6. Sincronizacion y autoridad.
    sync_rows = [
        ("Jugador propietario", "NetworkTransform con autoridad del propietario", "IMPLEMENTADO; TST-S3-006 pendiente"),
        ("Jugador propietario", "NetworkTransform replica la rotacion", "IMPLEMENTADO; TST-S3-007 pendiente"),
        ("Input local del propietario", "Move solo en IsOwner; transform replicado por NGO", "IMPLEMENTADO; TST-S3-008 pendiente"),
        ("Input local del propietario", "Salto local; posicion resultante replicada por NetworkTransform", "IMPLEMENTADO; TST-S3-009 pendiente"),
        ("Servidor valida; propietario aplica", "RespawnOwnerRpc y NetworkTransform.Teleport", "IMPLEMENTADO; TST-S3-010 pendiente"),
    ]
    for row, values in enumerate(sync_rows, start=1):
        cell(t[10], row, 1, values[0])
        cell(t[10], row, 2, values[1])
        cell(t[10], row, 4, values[2])
    set_labeled_paragraph(
        p[56],
        "Resultado general de sincronizacion: ",
        "IMPLEMENTADO / COMPILA. LE0-99 permanece en Pruebas y no existe una "
        "ejecucion manual identificable de TST-S3-006 a 010; resultado funcional pendiente."
    )

    # 7. Integracion con mecanicas previas.
    set_paragraph_text(
        p[59],
        "El networking integra el nucleo jugable y las mecanicas de parkour. La "
        "configuracion efectiva del prefab Jugador es velocidad 6, carrera 10, "
        "aceleracion 55, salto 7 y plano cercano de camara 0,05. El refinamiento "
        "de colision proyecta la velocidad solicitada sobre las paredes para "
        "conservar el movimiento tangencial, permitir alejarse y resolver esquinas "
        "sin modificar suelo, pendientes, techos ni WallRun."
    )
    cell(t[11], 0, 2, "Estado Jira 01/09")
    cell(t[11], 1, 2, "LE0-91 a 95 y LE0-113 finalizadas; LE0-100 en curso.")
    cell(t[11], 1, 3, "4 casos automatizados de respuesta contra paredes implementados y compilados; ejecucion y TST-S3-014 a 018/027 pendientes.")
    cell(t[11], 2, 2, "LE0-102, LE0-110 y LE0-113 finalizadas.")
    cell(t[11], 2, 3, "Implementado; TST-S3-019 a 021 y validacion multiplayer pendientes.")
    cell(t[11], 3, 2, "LE0-105, LE0-111 y LE0-113 finalizadas.")
    cell(t[11], 3, 3, "Implementado; prueba de integracion multiplayer pendiente.")
    cell(t[11], 4, 2, "LE0-103 en curso; LE0-113 finalizada.")
    cell(t[11], 4, 3, "TST-S3-022 a 024 pendientes.")
    cell(t[11], 5, 2, "LE0-104 en Pruebas; LE0-108 y LE0-113 finalizadas.")
    cell(t[11], 5, 3, "TST-S3-025 a 027 pendientes.")
    cell(t[11], 6, 2, "LE0-99 en Pruebas; LE0-113 finalizada; correcciones c33e5c4 y 99f19f3.")
    cell(t[11], 6, 3, "Implementado; TST-S3-010/024/028 con cliente remoto pendientes.")
    set_labeled_paragraph(
        p[61],
        "Resultado de integracion de principio a fin: ",
        "IMPLEMENTADO / COMPILA; recorrido completo con dos clientes pendiente. "
        "Baseline: rama feat/adicion-de-elementos-en-partida, commit c33e5c4."
    )

    # 8. Estado compartido y local.
    set_labeled_paragraph(
        p[63],
        "Trazabilidad Jira: ",
        "LE0-114 - Integracion de elementos en escenario y LE0-115 - "
        "Identificacion de estado de elementos del escenario; ambas finalizadas "
        "al corte del 01/09/2026."
    )
    set_paragraph_text(
        p[64],
        "La implementacion distingue el estado compartido del pickup y del efecto "
        "temporizado respecto del input, la camara y el feedback del propietario. "
        "El servidor valida la solicitud y decide un unico ganador; el estado se "
        "replica mediante NetworkObject y NetworkVariables. La decision queda "
        "documentada en ADR-S3-04."
    )
    cell(t[12], 1, 3, "CONFIRMADO: compartido mediante NetworkTransform.")
    cell(t[12], 2, 3, "CONFIRMADO: local y habilitado solo para IsOwner; near clip efectivo 0,05.")
    cell(t[12], 3, 3, "CONFIRMADO: servidor autoritativo y NetworkObject.Despawn(true).")
    cell(t[12], 4, 3, "CONFIRMADO: indicador y cuenta regresiva visibles solo para el propietario.")
    cell(t[12], 5, 1, "Compartido cuando afecta gameplay")
    cell(t[12], 5, 2, "El patron de movimiento y sus efectos deben converger entre clientes; LE0-103 sigue en curso.")
    cell(t[12], 5, 3, "PENDIENTE: validar autoridad y sincronizacion de obstaculos moviles.")
    set_labeled_paragraph(
        p[70],
        "Elemento elegido, owner y evidencia: ",
        "SpeedBoostPickup en GameScene. La configuracion serializada efectiva es "
        "x2 durante 10 s (prevalece sobre los valores por defecto x1,5/8 s del "
        "script). El servidor valida ownership y distancia, aplica el efecto con "
        "NetworkVariables y elimina el pickup para todos. El primer reclamo valido "
        "gana; la simultaneidad real queda pendiente de prueba manual."
    )

    # 9 y 10. Pruebas manuales multiplayer.
    for row in range(1, 8):
        cell(t[13], row, 2, "Unity 6000.3.22f1 / c33e5c4")
        cell(t[13], row, 3, "NO EJECUTADA")
        cell(t[13], row, 4, "Implementacion inspeccionada y compilada; sin evidencia manual de dos clientes.")
    cell(t[13], 3, 2, "Unity 6.3 LTS / 61dd628")
    cell(t[13], 3, 3, "PASS (registro QA)")
    cell(t[13], 3, 4, "Configuracion de escenas, NetworkManager, transporte, spawner y prefab validada por Sprint3ConfigurationTests.cs segun la hoja.")
    cell(t[13], 4, 2, "Unity 6.3 LTS / 61dd628")
    cell(t[13], 4, 3, "PASS (registro QA)")
    cell(t[13], 4, 4, "Prefab NetworkPlayer, Respawn y separacion de spawn validados por Sprint3ConfigurationTests.cs segun la hoja.")
    set_labeled_paragraph(
        p[76],
        "Resumen de la sesion: ",
        "el plan oficial contiene 30 casos. La hoja registra 2 ejecuciones automatizadas "
        "PASS (TST-S3-003 y 004), 0 FAIL y 28 casos no ejecutados: 6,7% del plan y "
        "100% de aprobacion sobre lo ejecutado. Las sesiones manuales multiplayer "
        "documentadas siguen en 0; los PASS no demuestran una conexion real de dos clientes."
    )
    for row in range(1, 6):
        cell(t[14], row, 3, "NO EJECUTADA. Implementacion disponible; evidencia con dos clientes pendiente.")
    for row in t[14].rows:
        prevent_row_split(row)

    # 11. Registro de problemas observados y soluciones.
    cell(t[15], 1, 1, "31/08")
    cell(t[15], 1, 2, "Respawn / sincronizacion")
    cell(t[15], 1, 3, "El cliente invitado no reaparecia tras caer; el Host sí.")
    cell(t[15], 1, 4, "2")
    cell(t[15], 1, 5, "Resuelto en codigo; regresion pendiente")
    cell(t[15], 1, 6, "99f19f3, c33e5c4; RespawnOwnerRpc.")
    cell(t[15], 2, 1, "31/08")
    cell(t[15], 2, 2, "Ownership / camara")
    cell(t[15], 2, 3, "Camara, audio o control podian activarse en una representacion no propietaria.")
    cell(t[15], 2, 4, "3")
    cell(t[15], 2, 5, "Resuelto en codigo; validacion pendiente")
    cell(t[15], 2, 6, "094ddd2; componentes solo para IsOwner.")
    cell(t[15], 3, 1, "01/09")
    cell(t[15], 3, 2, "Estado compartido")
    cell(t[15], 3, 3, "Reclamos simultaneos podian duplicar el power-up si cada cliente resolvia localmente.")
    cell(t[15], 3, 4, "3")
    cell(t[15], 3, 5, "Mitigado en codigo; prueba pendiente")
    cell(t[15], 3, 6, "36e0a04; servidor + ganador unico + Despawn(true).")

    # 12. Backlog actualizado.
    backlog_statuses = {
        1: ("Finalizada", ""),
        2: ("Finalizada", ""),
        3: ("Pruebas", ""),
        4: ("Finalizada", ""),
        5: ("Finalizada", "LE0-115 resuelta"),
    }
    for row, values in backlog_statuses.items():
        cell(t[16], row, 5, values[0])
        cell(t[16], row, 6, values[1])
    cell(t[16], 5, 4, "Bruno")

    transverse_statuses = {
        1: ("Finalizada", ""),
        2: ("Finalizada", "LE0-114 resuelta"),
        3: ("En curso", ""),
        4: ("En curso", ""),
        5: ("En curso", ""),
    }
    for row, values in transverse_statuses.items():
        cell(t[17], row, 5, values[0])
        cell(t[17], row, 6, values[1])

    # 13. Plan de pruebas y registro real de esta revision.
    set_paragraph_text(
        p[96],
        "La hoja Plan de Pruebas Definitivo - Sprint 3 es la fuente del catalogo "
        "oficial: 30 casos TST-S3, 16 criterios CA-S3 trazados, 26 casos de prioridad "
        "Alta y 18 automatizables. Registra PASS para TST-S3-003 y 004 en el commit "
        "61dd628 con evidencia Sprint3ConfigurationTests.cs. La fecha visible "
        "2026-01-09 es anterior al periodo del sprint y debe confirmarse; ademas, "
        "el commit y el archivo no existen en el checkout actual c33e5c4. Por eso "
        "se conserva el resultado reportado por la hoja, separado de la verificacion "
        "local y de las pruebas manuales aun pendientes."
    )
    cell(t[20], 1, 3, "13 planificados; TST-S3-003/004 PASS registrados; 11 pendientes")
    cell(t[20], 2, 3, "5 planificados; 4 tests locales apoyan TST-S3-017/018; no ejecutados")
    cell(t[20], 5, 3, "6 planificados; TST-S3-030 y pipeline CI pendientes")

    cell(t[21], 1, 0, "TST-S3-003")
    cell(t[21], 1, 1, "CA-S3-02/03 - LE0-31/97")
    cell(t[21], 1, 2, "61dd628 / 2026-01-09")
    cell(t[21], 1, 3, "PASS (hoja QA)")
    cell(t[21], 1, 4, "Configuracion de escenas, red, transporte, spawner y prefab validada.")
    cell(t[21], 1, 5, "Hoja QA; archivo/revision no disponibles en c33e5c4.")
    cell(t[21], 2, 0, "TST-S3-004")
    cell(t[21], 2, 1, "CA-S3-03 - LE0-97")
    cell(t[21], 2, 2, "61dd628 / 2026-01-09")
    cell(t[21], 2, 3, "PASS (hoja QA)")
    cell(t[21], 2, 4, "Prefab NetworkPlayer, Respawn y separacion de spawn validados.")
    cell(t[21], 2, 5, "Hoja QA; archivo/revision no disponibles en c33e5c4.")
    cell(t[21], 3, 0, "TST-S3-012/013\nTST-S3-017/018")
    cell(t[21], 3, 1, "CA-S3-06/10/13")
    cell(t[21], 3, 2, "c33e5c4 / revision local")
    cell(t[21], 3, 3, "PENDIENTE")
    cell(t[21], 3, 4, "Seis NUnit locales compilan: 4 de paredes y 2 de temporizador; no ejecutados.")
    cell(t[21], 3, 5, "WallCollisionResponseTests.cs; SpeedBoostTests.cs")

    # 14. Riesgos actualizados.
    risk_updates = {
        1: ("Mitigado / 1x3=3", "Relay y Host/Client integrados; versiones y flujo documentados. Sesion manual pendiente."),
        2: ("Abierto / 2x3=6", "LE0-99 en Pruebas; ejecutar TST-S3-006 a 010 con dos clientes."),
        3: ("Mitigado / 1x3=3", "Ownership local y autoridad del servidor definidos; ADR-S3-02; simultaneidad pendiente."),
        4: ("Mitigado / 1x2=2", "Versiones fijadas y compilacion correcta en Unity 6000.3.22f1."),
        5: ("Abierto / 3x3=9", "No hay sesion manual documentada; preparar Editor + build o dos builds."),
        6: ("Materializado / 3x3=9", "Solo 19/27 items y 78/118 SP finalizados al corte; priorizar alcance obligatorio."),
        7: ("En seguimiento / 2x2=4", "Se observan merges e integracion frecuente; mantener ramas cortas y PR."),
        8: ("En seguimiento / 2x2=4", "Codigo y ADR reducen dependencia, pero falta rotar ejecucion de pruebas multiplayer."),
        9: ("Materializado / 2x3=6", "La integracion existe, pero varias validaciones quedaron para el cierre."),
        10: ("Abierto / 2x2=4", "La hoja registra evidencia para 2/30 casos; faltan 28 resultados y la validacion manual."),
    }
    for row, values in risk_updates.items():
        cell(t[22], row, 3, values[0])
        cell(t[22], row, 4, values[1])

    cell(t[23], 1, 1, "Divergencia entre valores por defecto del script y valores serializados en escena/prefab.")
    cell(t[23], 1, 2, "2")
    cell(t[23], 1, 3, "2")
    cell(t[23], 1, 4, "4")
    cell(t[23], 1, 5, "Documentar valores efectivos y verificar escena antes de cada prueba. Owner: Tomas/Bruno.")
    cell(t[23], 1, 6, "GameScene: boost x2/10 s; Jugador.prefab: near clip 0,05 y aceleracion 55.")
    cell(t[23], 2, 1, "Ausencia de evidencia manual con dos clientes al cierre de la actualizacion.")
    cell(t[23], 2, 2, "3")
    cell(t[23], 2, 3, "3")
    cell(t[23], 2, 4, "9")
    cell(t[23], 2, 5, "Ejecutar matriz TST-S3-001 a 013 y 027; conservar capturas/logs. Owner: Testing/Tomas.")
    cell(t[23], 2, 6, "Registro actual: 0 sesiones manuales documentadas.")

    # 15. Metricas.
    cell(t[24], 1, 3, "70,37% por items (19/27). Complemento: 66,10% por SP (78/118).")
    cell(t[24], 2, 3, "Trazabilidad: 16/16 CA cubiertos. Validacion funcional: 2/30 casos ejecutados (6,7%).")
    cell(t[24], 3, 3, "0 al corte en Jira; LE0-115 ya no esta bloqueada.")
    cell(t[24], 4, 3, "3: T7, G1 y R-S3-02 con prioridad 9.")
    cell(t[24], 5, 3, "100% sobre ejecutadas (2/2); 6,7% sobre el plan total (2/30).")
    cell(t[24], 6, 3, "0 bugs abiertos en Jira; validaciones pendientes.")
    cell(t[24], 7, 3, "Sin datos: 0 sesiones documentadas.")
    cell(t[24], 8, 3, "Sin datos: TST-S3-006 a 010 pendientes.")
    set_paragraph_text(
        p[112],
        "Corte Jira del 01/09/2026: 27 items y 118 Story Points; 19 items "
        "(78 SP) finalizados, 7 items (38 SP) en curso o Pruebas y 1 item "
        "(2 SP) por hacer. Cumplimiento: 70,37% por items y 66,10% por SP. "
        "Plan QA: 30 casos, 16 criterios trazados, 2 PASS registrados y 28 pendientes."
    )
    p[108].paragraph_format.page_break_before = True

    # 16. ADRs.
    cell(t[25], 1, 1, "Unity 6000.3.22f1 (Unity 6.3 LTS)")
    cell(t[25], 6, 1, "Unity Test Framework 1.6.0 + GitHub Actions/GameCI; pruebas manuales multiplayer pendientes")

    decision_statuses = {
        1: "Adoptada; sesion manual pendiente.",
        2: "Adoptada; LE0-99 en Pruebas.",
        3: "Adoptada; TST-S3-006 a 010 pendientes.",
        4: "Adoptada; simultaneidad pendiente.",
    }
    for row, status in decision_statuses.items():
        cell(t[26], row, 3, status)

    cell(t[27], 1, 0, "ADR-S3-01 / 01-09-2026")
    cell(t[27], 1, 1, "Se comparo conexion directa, servidor dedicado y Relay Host/Client. Se priorizo alcance academico, baja infraestructura y soporte entre redes.")
    cell(t[27], 1, 2, "El Host crea allocation y join code de Relay; el cliente ingresa con ese codigo; el Host carga GameScene.")
    cell(t[27], 1, 3, "Simplifica despliegue y evita exponer IP. La sesion depende del Host y requiere prueba con dos instancias.")
    cell(t[27], 1, 4, "LE0-31; ConnectionManager.cs; commit 268b896; prueba manual pendiente.")
    cell(t[27], 2, 0, "ADR-S3-04 / 01-09-2026")
    cell(t[27], 2, 1, "Resolver el pickup localmente podia duplicar la obtencion. Se comparo estado local, autoridad del propietario y autoridad del servidor.")
    cell(t[27], 2, 2, "El servidor valida owner/distancia, aplica boost x2 por 10 s y despawnea el NetworkObject; el feedback queda local.")
    cell(t[27], 2, 3, "Un unico reclamo valido gana y todos ven la desaparicion. Queda pendiente probar solicitudes simultaneas reales.")
    cell(t[27], 2, 4, "LE0-114/115; SpeedBoostPickup.cs; SpeedBoostController.cs; GameScene; commit 36e0a04.")
    for row in t[27].rows:
        prevent_row_split(row)

    set_paragraph_text(p[122], "Detalle - ADR-S3-04: estado compartido del power-up")
    cell(t[28], 1, 1, "01/09/2026 - Tomas Molina Varas y Bruno Masdeu.")
    cell(t[28], 2, 1, "Un power-up exclusivo debe desaparecer para todos y no puede ser obtenido por dos jugadores ante reclamos concurrentes.")
    cell(t[28], 3, 1, "Resolucion local; autoridad del propietario; autoridad del servidor con validacion de owner y distancia.")
    cell(t[28], 4, 1, "Autoridad del servidor. Configuracion efectiva en GameScene: multiplicador x2 durante 10 s; Despawn(true) tras el primer reclamo valido.")
    cell(t[28], 5, 1, "El servidor procesa solicitudes secuencialmente, replica el temporizador y evita estados divergentes. Los valores serializados de escena son la fuente efectiva.")
    cell(t[28], 6, 1, "Beneficios: consistencia y ganador unico. Costos/riesgos: dependencia del Host y latencia; requiere validacion simultanea con dos clientes.")
    cell(t[28], 7, 1, "Jira: LE0-114/115. Pruebas: SpeedBoostTests.cs compila, no ejecutado; TST-S3-012/013 pendientes. Commit: 36e0a04.")

    # Checklist y fuentes.
    checklist = {
        125: "COMPLETADO: estado y fotografia del Sprint 3 actualizados desde Jira al 01/09/2026.",
        126: "PARCIAL: la hoja QA registra PASS para TST-S3-003/004; faltan 28 casos, pruebas manuales y evidencia verificable en el checkout actual.",
        127: "COMPLETADO: problemas reproducibles vinculados con codigo y commits; regresiones manuales identificadas.",
        128: "COMPLETADO: riesgos actualizados; tres riesgos criticos permanecen abiertos/materializados.",
        129: "COMPLETADO CON LIMITACIONES: metricas Jira y QA incorporadas; fecha/commit de los dos PASS requieren confirmacion local.",
        130: "COMPLETADO: ADR-S3-01 a 04 registradas y separadas de la linea de base heredada.",
        131: "COMPLETADO: los dieciseis entregables tienen evidencia disponible o una justificacion explicita de pendiente.",
    }
    for index, text in checklist.items():
        set_paragraph_text(p[index], text)

    set_paragraph_text(
        p[132],
        "Fuentes: Enunciado TP Inicial 2C 2026; Jira del proyecto LCS Equipo 07 "
        "consultado el 01/09/2026; documentos de Sprint 1 y Sprint 2; repositorio "
        "Unity en rama feat/adicion-de-elementos-en-partida, commit c33e5c4; "
        "historial de la tarea Refina la colision con paredes; Google Sheets "
        "Plan de Pruebas Definitivo - Sprint 3 (ID 1WRT_nUyYJo7DP44lCkpF3zIXL_IOmNxL1w91UeiCCt0), "
        "consultado el 01/09/2026."
    )

    # Dos saltos heredados generaban paginas casi vacias despues de ampliar el
    # contenido. El flujo natural conserva la jerarquia sin esos huecos.
    remove_page_break(p[46])
    remove_page_break(p[71])

    restore_spanish_diacritics(doc)

    # Metadatos de la copia.
    doc.core_properties.title = "Documentacion Sprint 3 - Integracion Multiplayer, Experimentacion y Testing"
    doc.core_properties.subject = "Sprint 3 actualizado al 01/09/2026"
    doc.core_properties.author = "LCS Equipo 07"
    doc.core_properties.last_modified_by = "LCS Equipo 07"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
